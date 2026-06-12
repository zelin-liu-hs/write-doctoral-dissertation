#!/usr/bin/env python3
"""Convert a Markdown dissertation draft to a readable DOCX file."""

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for kind, text in (
        ("begin", None),
        (None, "PAGE"),
        ("separate", None),
        (None, "1"),
        ("end", None),
    ):
        if kind:
            element = OxmlElement("w:fldChar")
            element.set(qn("w:fldCharType"), kind)
        else:
            element = OxmlElement("w:instrText" if text == "PAGE" else "w:t")
            if text == "PAGE":
                element.set(qn("xml:space"), "preserve")
            element.text = text
        run._r.append(element)


def add_inline_markdown(paragraph, text: str, base_font: str = "宋体", size: float = 11) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, base_font, size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, base_font, size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Consolas", size - 0.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, base_font, size)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(5)

    for level, size in ((1, 18), (2, 15), (3, 13), (4, 12)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "黑体"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    add_page_number(section.footer.paragraphs[0])


def add_table(document: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            add_inline_markdown(
                paragraph,
                row[column_index] if column_index < len(row) else "",
                size=9.5,
            )
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph()


def convert(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8-sig").splitlines()
    document = Document()
    configure_document(document)

    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("|") and "|" in line[1:]:
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue

        if not line:
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(
                paragraph,
                heading.group(2).strip(),
                base_font="黑体",
                size={1: 18, 2: 15, 3: 13, 4: 12}[level],
            )
            if first_heading:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(14)
                first_heading = False
            index += 1
            continue

        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.8)
            paragraph.paragraph_format.right_indent = Cm(0.5)
            add_inline_markdown(paragraph, line[2:].strip(), size=10.5)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(70, 70, 70)
            index += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            text = bullet.group(1).replace("[ ]", "□").replace("[x]", "■")
            add_inline_markdown(paragraph, text)
            index += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, numbered.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        add_inline_markdown(paragraph, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    with zipfile.ZipFile(output_path) as archive:
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError("DOCX validation failed: word/document.xml is missing")
    Document(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="UTF-8 Markdown input")
    parser.add_argument("output", type=Path, help="DOCX output")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if not args.input.is_file():
        raise FileNotFoundError(f"Markdown input does not exist: {args.input}")
    convert(args.input.resolve(), args.output.resolve())
    print(f"{args.output.resolve()}\t{args.output.resolve().stat().st_size}")


if __name__ == "__main__":
    main()
